from docx import Document
from docx.shared import Pt
from fpdf import FPDF
from datetime import datetime

def export_to_docx(filename, content_dict, title="Report", author="Analyst"):
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(f"Author: {author}")
    doc.add_paragraph(f"Date: {datetime.today().strftime('%Y-%m-%d')}")
    doc.add_page_break()

    for section_title, section_body in content_dict.items():
        doc.add_heading(section_title, level=1)
        p = doc.add_paragraph(section_body)
        p.style.font.size = Pt(11)
        doc.add_paragraph("\n")

    doc.save(filename)
    return filename


def export_to_pdf(filename, content_dict, title="Report", author="Analyst"):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", 'B', 20)
    pdf.cell(0, 10, title, ln=True, align="C")
    pdf.set_font("Arial", '', 12)
    pdf.cell(0, 10, f"Author: {author}", ln=True, align="C")
    pdf.cell(0, 10, f"Date: {datetime.today().strftime('%Y-%m-%d')}", ln=True, align="C")
    pdf.add_page()

    for section_title, section_body in content_dict.items():
        pdf.set_font("Arial", 'B', 14)
        pdf.multi_cell(0, 10, section_title)
        pdf.set_font("Arial", '', 12)
        pdf.multi_cell(0, 10, section_body)
        pdf.ln(5)

    pdf.output(filename)
    return filename
