import streamlit as st
import pandas as pd
from PyPDF2 import PdfReader
import docx
from docx import Document
from fpdf import FPDF
import openai
import os
from datetime import datetime

# --- Page Config ---
st.set_page_config(
    page_title="Generative AI Portfolio",
    page_icon="🤖",
    layout="wide",
    initial_sidebar_state="expanded"
)

# --- Load OpenAI API key ---
openai.api_key = os.getenv("OPENAI_API_KEY")


# --- Utility Functions ---

def call_openai(prompt, max_tokens=300):
    try:
        response = openai.ChatCompletion.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are a helpful assistant."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=max_tokens
        )
        return response.choices[0].message["content"].strip()
    except Exception as e:
        return f"⚠️ Error: {e}"


def transcribe_audio(file):
    try:
        audio_file = open(file, "rb")
        transcript = openai.Audio.transcriptions.create(model="whisper-1", file=audio_file)
        return transcript.text
    except Exception as e:
        return f"⚠️ Error: {e}"


def extract_text_from_file(uploaded_file):
    text = ""
    if uploaded_file.name.endswith(".pdf"):
        reader = PdfReader(uploaded_file)
        for page in reader.pages:
            text += page.extract_text()
    elif uploaded_file.name.endswith(".docx"):
        doc = docx.Document(uploaded_file)
        for para in doc.paragraphs:
            text += para.text + "\n"
    elif uploaded_file.name.endswith(".xlsx"):
        df = pd.read_excel(uploaded_file)
        text = df.to_string()
    return text


def clean_text(text):
    # Remove emojis or non-latin characters for PDF export
    return text.encode("latin-1", errors="ignore").decode("latin-1")


# Word Export
def export_to_docx(filename, content_dict, title="Report", author="Analyst", logo_path=None):
    doc = Document()
    # Cover Page
    if logo_path:
        doc.add_picture(logo_path, width=None)
    doc.add_heading(title, 0)
    doc.add_paragraph(f"Author: {author}")
    doc.add_paragraph(f"Date: {datetime.today().strftime('%Y-%m-%d')}")
    doc.add_page_break()
    # Content Sections
    for section_title, section_body in content_dict.items():
        doc.add_heading(section_title, level=1)
        doc.add_paragraph(section_body)
    doc.save(filename)
    return filename


# PDF Export
def export_to_pdf(filename, content_dict, title="Report", author="Analyst", logo_path=None):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", 'B', 16)
    pdf.cell(0, 10, clean_text(title), ln=True, align="C")
    pdf.set_font("Arial", '', 12)
    pdf.cell(0, 10, f"Author: {clean_text(author)}", ln=True, align="C")
    pdf.cell(0, 10, f"Date: {datetime.today().strftime('%Y-%m-%d')}", ln=True, align="C")
    pdf.ln(10)
    for section_title, section_body in content_dict.items():
        pdf.set_font("Arial", 'B', 14)
        pdf.multi_cell(0, 10, clean_text(section_title))
        pdf.set_font("Arial", '', 12)
        pdf.multi_cell(0, 10, clean_text(section_body))
        pdf.ln(5)
    pdf.output(filename)
    return filename


# --- Sidebar ---
st.sidebar.title("🚀 GenAI Portfolio")
app_mode = st.sidebar.radio(
    "Choose App",
    [
        "🏢 Meeting Intelligence Assistant",
        "📑 Requirement → User Story Translator",
        "📊 Customer Feedback Analyzer",
        "⚖ Regulatory Change Summarizer"
    ]
)
st.sidebar.markdown("---")
st.sidebar.info("Built with Streamlit + OpenAI\nAuthor: Bhagyashree Deshmukh")

# --- Main ---
st.title("🤖 Generative AI Portfolio")
st.write("Showcasing practical AI use cases for Business Analysis & Product Ownership.")

# --- App: Meeting Intelligence Assistant ---
if app_mode == "🏢 Meeting Intelligence Assistant":
    st.header("📋 Meeting Intelligence Assistant")
    uploaded_audio = st.file_uploader("Upload meeting audio (mp3, wav, m4a)", type=["mp3", "wav", "m4a"])
    if uploaded_audio:
        progress = st.progress(0)
        status = st.empty()
        # Step 1: Transcription
        status.text("Step 1/3: Transcribing audio...")
        text = transcribe_audio(uploaded_audio)
        progress.progress(33)
        # Step 2: Summarization
        status.text("Step 2/3: Summarizing meeting...")
        summary = call_openai(f"Summarize this meeting:\n\n{text}", max_tokens=200)
        progress.progress(66)
        # Step 3: Action items
        status.text("Step 3/3: Extracting action items...")
        actions = call_openai(f"Extract action items from this meeting:\n\n{text}", max_tokens=150)
        progress.progress(100)
        status.text("✅ Done! Meeting processed successfully.")
        st.subheader("✨ Summary");
        st.write(summary)
        st.subheader("📝 Action Items");
        st.write(actions)
        with st.expander("📄 Transcript"): st.write(text)
        # Export
        report_content = {"Summary": summary, "Action Items": actions, "Transcript": text}
        docx_file = "meeting_report.docx"
        export_to_docx(docx_file, report_content, title="Meeting Report")
        with open(docx_file, "rb") as f:
            st.download_button("⬇️ Download Word Report", f, file_name=docx_file,
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        pdf_file = "meeting_report.pdf"
        export_to_pdf(pdf_file, report_content, title="Meeting Report")
        with open(pdf_file, "rb") as f:
            st.download_button("⬇️ Download PDF Report", f, file_name=pdf_file, mime="application/pdf")

# --- App: Requirement → User Story Translator ---
elif app_mode == "📑 Requirement → User Story Translator":
    st.header("📑 Requirement → User Story Translator")
    uploaded_req = st.file_uploader("Upload requirements (PDF, Word, Excel)", type=["pdf", "docx", "xlsx"])
    if uploaded_req:
        progress = st.progress(0)
        status = st.empty()
        status.text("Step 1/2: Extracting text from document...")
        text = extract_text_from_file(uploaded_req)
        progress.progress(50)
        status.text("Step 2/2: Translating into user stories...")
        stories = call_openai(
            f"Convert the following requirements into Agile user stories with acceptance criteria:\n\n{text}",
            max_tokens=300)
        progress.progress(100)
        status.text("✅ Done! Requirements converted into user stories.")
        st.subheader("📌 User Stories");
        st.write(stories)
        report_content = {"User Stories": stories}
        docx_file = "user_stories.docx"
        export_to_docx(docx_file, report_content, title="User Story Report")
        with open(docx_file, "rb") as f:
            st.download_button("⬇️ Download Word Report", f, file_name=docx_file,
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        pdf_file = "user_stories.pdf"
        export_to_pdf(pdf_file, report_content, title="User Story Report")
        with open(pdf_file, "rb") as f:
            st.download_button("⬇️ Download PDF Report", f, file_name=pdf_file, mime="application/pdf")

# --- App: Customer Feedback Analyzer ---
elif app_mode == "📊 Customer Feedback Analyzer":
    st.header("📊 Customer Feedback Analyzer")
    uploaded_csv = st.file_uploader("Upload CSV with 'feedback' column", type=["csv"])
    if uploaded_csv:
        df = pd.read_csv(uploaded_csv)
        st.write("📄 Sample Data", df.head())
        progress = st.progress(0)
        status = st.empty()
        status.text("Step 1/3: Preparing feedback data...")
        all_feedback = " ".join(df["feedback"].astype(str).tolist())
        progress.progress(33)
        status.text("Step 2/3: Sentiment analysis...")
        sentiment = call_openai(f"Analyze sentiment distribution of this feedback:\n\n{all_feedback}")
        progress.progress(66)
        status.text("Step 3/3: Identifying key themes...")
        themes = call_openai(f"Identify top 3 recurring themes from this feedback:\n\n{all_feedback}")
        progress.progress(100)
        status.text("✅ Done! Feedback analysis complete.")
        st.subheader("📊 Sentiment Analysis");
        st.write(sentiment)
        st.subheader("✨ Key Themes");
        st.write(themes)
        # Export
        report_content = {"Sentiment Analysis": sentiment, "Key Themes": themes}
        docx_file = "customer_feedback.docx"
        export_to_docx(docx_file, report_content, title="Customer Feedback Analysis")
        with open(docx_file, "rb") as f:
            st.download_button("⬇️ Download Word Report", f, file_name=docx_file,
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        pdf_file = "customer_feedback.pdf"
        export_to_pdf(pdf_file, report_content, title="Customer Feedback Analysis")
        with open(pdf_file, "rb") as f:
            st.download_button("⬇️ Download PDF Report", f, file_name=pdf_file, mime="application/pdf")

# --- App: Regulatory Change Summarizer ---
elif app_mode == "⚖ Regulatory Change Summarizer":
    st.header("⚖ Regulatory Change Summarizer")
    uploaded_pdf = st.file_uploader("Upload regulatory PDF", type=["pdf"])
    if uploaded_pdf:
        progress = st.progress(0)
        status = st.empty()
        status.text("Step 1/3: Extracting regulation text...")
        reader = PdfReader(uploaded_pdf)

