from utils.export_utils import export_to_docx, export_to_pdf
from utils.file_utils import extract_text_from_file
import streamlit as st
import pandas as pd
import altair as alt
from PyPDF2 import PdfReader
import openai
import os

from docx import Document
from fpdf import FPDF

def export_to_docx(filename, content, title="Report"):
    """Save content to a Word document"""
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(content)
    doc.save(filename)
    return filename

def export_to_pdf(filename, content, title="Report"):
    """Save content to a PDF file"""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", 'B', 16)
    pdf.multi_cell(0, 10, title, align="C")
    pdf.ln(10)
    pdf.set_font("Arial", '', 12)
    pdf.multi_cell(0, 10, content)
    pdf.output(filename)
    return filename

from docx.shared import Pt
from datetime import datetime

def export_to_docx(filename, content_dict, title="Report", author="Analyst", logo_path=None):
    """Save structured content to a Word document with a cover page"""
    doc = Document()

    # --- Cover Page ---
    if logo_path:
        doc.add_picture(logo_path, width=None)  # auto-scale logo
    doc.add_heading(title, 0)
    doc.add_paragraph(f"Author: {author}")
    doc.add_paragraph(f"Date: {datetime.today().strftime('%Y-%m-%d')}")
    doc.add_page_break()

    # --- Content Sections ---
    for section_title, section_body in content_dict.items():
        doc.add_heading(section_title, level=1)
        p = doc.add_paragraph(section_body)
        p.style.font.size = Pt(11)
        doc.add_paragraph("\n")

    doc.save(filename)
    return filename


def export_to_pdf(filename, content_dict, title="Report", author="Analyst", logo_path=None):
    """Save structured content to a PDF with a cover page"""
    pdf = FPDF()
    pdf.add_page()

    # --- Cover Page ---
    if logo_path:
        pdf.image(logo_path, x=60, y=20, w=90)  # centered
        pdf.ln(60)
    pdf.set_font("Arial", 'B', 20)
    pdf.cell(0, 10, title, ln=True, align="C")
    pdf.set_font("Arial", '', 12)
    pdf.cell(0, 10, f"Author: {author}", ln=True, align="C")
    pdf.cell(0, 10, f"Date: {datetime.today().strftime('%Y-%m-%d')}", ln=True, align="C")
    pdf.add_page()

    # --- Content Sections ---
    pdf.set_font("Arial", '', 12)
    for section_title, section_body in content_dict.items():
        pdf.set_font("Arial", 'B', 14)
        pdf.multi_cell(0, 10, section_title)
        pdf.set_font("Arial", '', 12)
        pdf.multi_cell(0, 10, section_body)
        pdf.ln(5)

    pdf.output(filename)
    return filename


# --- Page Config ---
st.set_page_config(
    page_title="GenAI Portfolio",
    page_icon="🤖",
    layout="wide",
    initial_sidebar_state="expanded",
primaryColor = "#4B9CD3",
backgroundColor = "#F8F9FA",
secondaryBackgroundColor = "#FFFFFF",
textColor = "#333333",
font = "sans serif"

)




# --- Load API Key ---
openai.api_key = os.getenv("OPENAI_API_KEY")

def call_openai(prompt, max_tokens=300):
    """Helper function to call OpenAI GPT model"""
    try:
        response = openai.ChatCompletion.create(
            model="gpt-4o-mini",  # use gpt-3.5-turbo if quota limited
            messages=[{"role": "system", "content": "You are a helpful assistant."},
                      {"role": "user", "content": prompt}],
            max_tokens=max_tokens
        )
        return response.choices[0].message["content"].strip()
    except Exception as e:
        return f"⚠️ Error: {e}"

# --- Sidebar ---
st.sidebar.title("🚀 GenAI Portfolio")
app_mode = st.sidebar.radio(
    "Choose App",
    [
        "🏢 Meeting Intelligence Assistant",
        "📑 Requirement → User Story Translator",
        "📊 Customer Feedback Analyzer",
        "⚖ Regulatory Change Summarizer"
    ]
)
st.sidebar.markdown("---")
st.sidebar.info("Built with Streamlit + OpenAI\n\n**Author:** Bhagyashree Deshmukh")

# --- Main Title ---
st.title("🤖 Generative AI Portfolio")
st.write("Showcasing practical AI use cases for Business Analysis & Product Ownership.")

# --- Meeting Intelligence Assistant ---
if app_mode == "🏢 Meeting Intelligence Assistant":
    st.header("📋 Meeting Intelligence Assistant")
    uploaded_audio = st.file_uploader("Upload meeting audio (mp3, wav, m4a)", type=["mp3", "wav", "m4a"])

    if uploaded_audio:
        progress = st.progress(0)
        status = st.empty()

        # Step 1: Transcription
        status.text("Step 1/3: Transcribing audio...")
        text = transcribe_audio(uploaded_audio)
        progress.progress(33)

        # Step 2: Summarization
        status.text("Step 2/3: Summarizing meeting...")
        summary = call_openai(f"Summarize this meeting:\n\n{text}", max_tokens=200)
        progress.progress(66)

        # Step 3: Action items
        status.text("Step 3/3: Extracting action items...")
        actions = call_openai(f"Extract action items from this meeting:\n\n{text}", max_tokens=150)
        progress.progress(100)

        status.text("✅ Done! Meeting processed successfully.")

        st.subheader("✨ Summary")
        st.write(summary)
        st.subheader("📝 Action Items")
        st.write(actions)

        with st.expander("📄 Transcript"):
            st.write(text)

        # --- Export ---
        #export_text = f"📋 Meeting Summary\n\n{summary}\n\n📝 Action Items\n\n{actions}\n\n---\nTranscript:\n{text}"
report_content = {
            "📋 Meeting Summary": summary,
            "📝 Action Items": actions,
            "📄 Transcript": text
}
        # Export DOCX
docx_file = "meeting_report.docx"
export_to_docx(docx_file, report_content, title="Meeting Report", author="Bhagyashree Deshmukh")
with open(docx_file, "rb") as f:
            st.download_button("⬇️ Download Word Report", f, file_name=docx_file, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        # Export PDF
pdf_file = "meeting_report.pdf"
export_to_pdf(pdf_file, report_content, title="Meeting Report", author="Bhagyashree Deshmukh")
with open(pdf_file, "rb") as f:
            st.download_button("⬇️ Download PDF Report", f, file_name=pdf_file, mime="application/pdf")

# --- Requirement → User Story Translator ---

elif app_mode == "📑 Requirement → User Story Translator":
    st.header("📑 Requirement → User Story Translator")
    uploaded_req = st.file_uploader("Upload requirements (PDF, Word, or Excel)", type=["pdf", "docx", "xlsx"])

    if uploaded_req:
        progress = st.progress(0)
        status = st.empty()

        # Step 1: Extract text
        status.text("Step 1/2: Extracting text from document...")
        text = extract_text_from_file(uploaded_req)
        progress.progress(50)

        # Step 2: Translate into user stories
        status.text("Step 2/2: Translating requirements into user stories...")
        stories = call_openai(
            f"Convert the following requirements into Agile user stories with acceptance criteria:\n\n{text}",
            max_tokens=300
        )
        progress.progress(100)

        status.text("✅ Done! Requirements converted into user stories.")

        st.subheader("📌 User Stories")
        st.write(stories)

        # --- Export ---
        #export_text = f"📑 User Stories\n\n{stories}"
report_content = {
            "📑 User Stories": stories
        }
        # Export DOCX
        docx_file = "user_stories.docx"
        export_to_docx(docx_file, report_content, title="User Stories",author="Bhagyashree Deshmukh")
        with open(docx_file, "rb") as f:
            st.download_button("⬇️ Download Word User Stories", f, file_name=docx_file, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        # Export PDF
        pdf_file = "user_stories.pdf"
        export_to_pdf(pdf_file, report_content, title="User Stories",author="Bhagyashree Deshmukh")
        with open(pdf_file, "rb") as f:
            st.download_button("⬇️ Download PDF User Stories", f, file_name=pdf_file, mime="application/pdf")

# --- Customer Feedback Analyzer ---
elif app_mode == "📊 Customer Feedback Analyzer":
    st.header("📊 Customer Feedback Analyzer")
    uploaded_csv = st.file_uploader("Upload CSV (must contain a 'feedback' column)", type=["csv"])

    if uploaded_csv:
        df = pd.read_csv(uploaded_csv)
        st.write("📄 Sample Data", df.head())

        progress = st.progress(0)
        status = st.empty()

        # Step 1: Combine all feedback
        status.text("Step 1/3: Preparing feedback data...")
        all_feedback = " ".join(df["feedback"].astype(str).tolist())
        progress.progress(33)

        # Step 2: Sentiment analysis
        status.text("Step 2/3: Analyzing sentiment distribution...")
        sentiment = call_openai(f"Analyze sentiment distribution (Positive, Neutral, Negative) of this feedback:\n\n{all_feedback}")
        progress.progress(66)

        # Step 3: Identify themes
        status.text("Step 3/3: Identifying key themes...")
        themes = call_openai(f"Identify top 3 recurring themes in this customer feedback:\n\n{all_feedback}")
        progress.progress(100)

        status.text("✅ Done! Feedback analysis complete.")

        # --- Display Results ---
        st.subheader("📊 Sentiment Analysis")
        st.write(sentiment)

        st.subheader("✨ Key Themes")
        st.write(themes)

        # Optional: Dummy chart (replace with actual sentiment extraction later)
        chart_data = pd.DataFrame({
            "Sentiment": ["Positive", "Neutral", "Negative"],
            "Count": [45, 20, 10]
        })
        chart = alt.Chart(chart_data).mark_bar().encode(x="Sentiment", y="Count", color="Sentiment")
        st.altair_chart(chart, use_container_width=True)

        # --- Export ---
        #export_text = f"📊 Sentiment Analysis\n\n{sentiment}\n\n✨ Key Themes\n\n{themes}"
report_content = {
            "📊 Sentiment Analysis": sentiment,
            "✨ Key Themes": themes
        }
        
        # Export DOCX
        docx_file = "customer_feedback.docx"
        export_to_docx(docx_file, export_text, title="Customer Feedback Analysis",author="Bhagyashree Deshmukh")
        with open(docx_file, "rb") as f:
            st.download_button("⬇️ Download Word Report", f, file_name=docx_file, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        # Export PDF
        pdf_file = "customer_feedback.pdf"
        export_to_pdf(pdf_file, export_text, title="Customer Feedback Analysis",author="Bhagyashree Deshmukh")
        with open(pdf_file, "rb") as f:
            st.download_button("⬇️ Download PDF Report", f, file_name=pdf_file, mime="application/pdf")

# --- Regulatory Change Summarizer ---
elif app_mode == "⚖ Regulatory Change Summarizer":
    st.header("⚖ Regulatory Change Summarizer")
    uploaded_pdf = st.file_uploader("Upload regulatory PDF", type=["pdf"])

    if uploaded_pdf:
        progress = st.progress(0)
        status = st.empty()

        # Step 1: Extract text
        status.text("Step 1/3: Extracting regulation text...")
        reader = PdfReader(uploaded_pdf)
        text = ""
        for page in reader.pages:
            text += page.extract_text()
        progress.progress(33)

        # Step 2: Summarize regulation
        status.text("Step 2/3: Summarizing regulation...")
        summary = call_openai(f"Summarize this regulation in simple business terms:\n\n{text}", max_tokens=250)
        progress.progress(66)

        # Step 3: Business impact analysis
        status.text("Step 3/3: Identifying business impacts...")
        impact = call_openai(f"What are the business and compliance impacts of this regulation?\n\n{text}", max_tokens=250)
        progress.progress(100)

        status.text("✅ Done! Regulation analyzed successfully.")

        # --- Display Results ---
        st.subheader("📌 Summary")
        st.write(summary)

        st.subheader("💡 Business Impact")
        st.write(impact)

        # --- Export ---
        #export_text = f"⚖ Regulatory Change Summary\n\n📌 Summary\n{summary}\n\n💡 Business Impact\n{impact}"
report_content = {
            "📌 Summary": summary,
            "💡 Business Impact": impact
        }
        # Export DOCX
        docx_file = "regulatory_summary.docx"
        export_to_docx(docx_file, report_content, title="Regulatory Change Summary",author="Bhagyashree Deshmukh")
        with open(docx_file, "rb") as f:
            st.download_button("⬇️ Download Word Report", f, file_name=docx_file, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        # Export PDF
        pdf_file = "regulatory_summary.pdf"
        export_to_pdf(pdf_file, report_content, title="Regulatory Change Summary",author="Bhagyashree Deshmukh")
        with open(pdf_file, "rb") as f:
            st.download_button("⬇️ Download PDF Report", f, file_name=pdf_file, mime="application/pdf")
